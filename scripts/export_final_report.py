#!/usr/bin/env python3
"""Build the final safety/non-safety wage-mechanism report.

This script is intentionally local-only: it reads Markdown and image assets
from the repository, then writes DOCX/PDF artifacts with python-docx and
ReportLab. It does not call GABRIEL, external models, APIs, or the network.
"""

from __future__ import annotations

import argparse
import csv
import os
import re
import shutil
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


REPORT_TITLE = "Deeper Look Into Safety & Non-safety Wage Mechanisms"
REPORT_SUBTITLE = "Mechanisms, Counterpoints, and Source Needs Across Municipal Occupations"
REPORT_LABEL = "Safety & Non-safety Wage Mechanisms"
REPORT_DATE = "July 10, 2026"
CRIMSON = "A51C30"
CHARCOAL = "333333"
MUTED = "666666"


@dataclass
class MarkdownBlock:
    kind: str
    text: str = ""
    level: int = 0
    alt: str = ""
    path: str = ""
    rows: list[list[str]] | None = None
    ordered: bool = False
    items: list[str] | None = None


def run(cmd: list[str]) -> str:
    proc = subprocess.run(cmd, text=True, capture_output=True, check=False)
    if proc.returncode != 0:
        raise RuntimeError(
            f"Command failed ({proc.returncode}): {' '.join(cmd)}\n"
            f"stdout:\n{proc.stdout}\nstderr:\n{proc.stderr}"
        )
    return proc.stdout.strip()


def integrate_markdown(scaffold: Path, appendix: Path, output: Path) -> None:
    scaffold_text = scaffold.read_text(encoding="utf-8")
    appendix_text = appendix.read_text(encoding="utf-8")

    scaffold_text = scaffold_text.replace(
        "**Status: DRAFT SCAFFOLD for human review.** Markdown, CSV tables, and charts only — no PDF/DOCX has been produced. This document is the input to a formatting run, not the final artifact.\n\n---\n\n",
        "",
    )
    scaffold_text = scaffold_text.replace(
        "**Status: DRAFT SCAFFOLD for human review.** Markdown, CSV tables, and charts only -- no PDF/DOCX has been produced. This document is the input to a formatting run, not the final artifact.\n\n---\n\n",
        "",
    )

    main, sep, _old_appendix = scaffold_text.partition("\n## Appendix\n")
    if not sep:
        raise ValueError("Could not find scaffold Appendix section to replace")

    appendix_body = re.sub(r"^# Report Appendix — 2026-07-10\s*\n+", "", appendix_text)
    appendix_body = re.sub(r"^# Report Appendix -- 2026-07-10\s*\n+", "", appendix_body)
    appendix_body = appendix_body.replace(
        "Supporting tables for `docs/analysis/report_scaffold_safety_non_safety_wage_mechanisms_2026-07-10.md`. Deterministic reference material only — no GABRIEL/codify, Harvard Proxy, or model/API calls were made to produce this appendix; every table below is compiled directly from `docs/analysis/gabriel_codify_evidence_layer.csv` and `docs/analysis/report_assets/`.",
        "Supporting tables for the integrated report. Deterministic reference material only -- no GABRIEL/codify, Harvard Proxy, or model/API calls were made to produce this appendix; every table below is compiled directly from `docs/analysis/gabriel_codify_evidence_layer.csv` and `docs/analysis/report_assets/`.",
    )
    integrated = main.rstrip() + "\n\n# Appendix\n\n" + appendix_body.lstrip()
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(integrated.rstrip() + "\n", encoding="utf-8")


def parse_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text


def parse_markdown(md_path: Path) -> list[MarkdownBlock]:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    blocks: list[MarkdownBlock] = []
    para: list[str] = []

    def flush_para() -> None:
        nonlocal para
        if para:
            blocks.append(MarkdownBlock("paragraph", " ".join(x.strip() for x in para).strip()))
            para = []

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            flush_para()
            i += 1
            continue
        if stripped == "---":
            flush_para()
            blocks.append(MarkdownBlock("rule"))
            i += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_para()
            blocks.append(MarkdownBlock("heading", parse_inline(heading.group(2)), len(heading.group(1))))
            i += 1
            continue
        image = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
        if image:
            flush_para()
            blocks.append(MarkdownBlock("image", alt=image.group(1), path=image.group(2)))
            i += 1
            continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_para()
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
                    rows.append([parse_inline(c) for c in cells])
                i += 1
            if rows:
                blocks.append(MarkdownBlock("table", rows=rows))
            continue
        bullet_items: list[str] = []
        ordered = False
        while i < len(lines):
            m = re.match(r"^\s*-\s+(.+)$", lines[i])
            n = re.match(r"^\s*\d+\.\s+(.+)$", lines[i])
            if m:
                bullet_items.append(parse_inline(m.group(1)))
                i += 1
                continue
            if n:
                ordered = True
                bullet_items.append(parse_inline(n.group(1)))
                i += 1
                continue
            break
        if bullet_items:
            flush_para()
            blocks.append(MarkdownBlock("list", ordered=ordered, items=bullet_items))
            continue
        para.append(stripped)
        i += 1
    flush_para()
    return blocks


def check_images(md_path: Path) -> list[Path]:
    text = md_path.read_text(encoding="utf-8")
    paths = [Path(p) for _alt, p in re.findall(r"!\[([^\]]*)\]\(([^)]+)\)", text)]
    missing = [p for p in paths if not (md_path.parent / p).exists()]
    if missing:
        raise FileNotFoundError("Missing image references: " + ", ".join(str(p) for p in missing))
    return [md_path.parent / p for p in paths]


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Georgia"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_name
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_docx_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for s in doc.sections:
        s.top_margin = Inches(0.82)
        s.bottom_margin = Inches(0.82)
        s.left_margin = Inches(0.82)
        s.right_margin = Inches(0.82)
        s.header_distance = Inches(0.35)
        s.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Georgia"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after in [
        ("Heading 1", 18, 18, 8),
        ("Heading 2", 14, 14, 6),
        ("Heading 3", 12, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Georgia"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(CHARCOAL if name != "Heading 2" else CRIMSON)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = REPORT_LABEL
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, 9, False, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Page ")
    set_run_font(r, 9, False, MUTED)
    add_field(footer, "PAGE")


def add_docx_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(REPORT_TITLE)
    set_run_font(r, 26, True, CHARCOAL)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(REPORT_SUBTITLE)
    set_run_font(r, 14, False, MUTED)
    p.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(REPORT_LABEL)
    set_run_font(r, 10.5, True, CRIMSON)
    p.paragraph_format.space_after = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(REPORT_DATE)
    set_run_font(r, 10.5, False, MUTED)
    doc.add_page_break()


def add_docx_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = True
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.cell(ri, ci)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[ci] if ci < len(row) else ""
            cell.text = text
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, 7 if cols > 5 else 8.5, ri == 0, "000000")
            if ri == 0:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F2F2F2")
                tc_pr.append(shd)
    doc.add_paragraph()


def build_docx(md_path: Path, out_path: Path) -> None:
    blocks = parse_markdown(md_path)
    doc = Document()
    set_docx_styles(doc)
    add_docx_cover(doc)

    skip_title_block = True
    for block in blocks:
        if skip_title_block:
            if block.kind == "heading" and block.level == 1 and block.text == REPORT_TITLE:
                continue
            if block.kind == "paragraph" and (
                block.text.startswith("Subtitle:")
                or block.text.startswith("Header/report label:")
                or block.text.startswith("**Subtitle:**")
                or block.text.startswith("**Header/report label:**")
            ):
                continue
            skip_title_block = False

        if block.kind == "heading":
            if block.text == "Appendix":
                doc.add_page_break()
            level = min(max(block.level, 1), 3)
            doc.add_heading(block.text, level=level)
        elif block.kind == "paragraph":
            p = doc.add_paragraph(parse_inline(block.text))
            p.paragraph_format.keep_together = False
        elif block.kind == "list":
            style = "List Number" if block.ordered else "List Bullet"
            for item in block.items or []:
                p = doc.add_paragraph(parse_inline(item), style=style)
                p.paragraph_format.space_after = Pt(4)
        elif block.kind == "image":
            image_path = md_path.parent / block.path
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(image_path), width=Inches(6.25))
            cap = doc.add_paragraph(block.alt)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                set_run_font(run, 9, False, MUTED)
        elif block.kind == "table":
            add_docx_table(doc, block.rows or [])
        elif block.kind == "rule":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p_border = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), CRIMSON)
            p_border.append(bottom)
            p._p.get_or_add_pPr().append(p_border)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def register_pdf_fonts() -> tuple[str, str]:
    candidates = [
        (
            "/System/Library/Fonts/Supplemental/Georgia.ttf",
            "/System/Library/Fonts/Supplemental/Georgia Bold.ttf",
        ),
        (
            "/Library/Fonts/Georgia.ttf",
            "/Library/Fonts/Georgia Bold.ttf",
        ),
    ]
    for regular, bold in candidates:
        if Path(regular).exists() and Path(bold).exists():
            pdfmetrics.registerFont(TTFont("GeorgiaLocal", regular))
            pdfmetrics.registerFont(TTFont("GeorgiaLocal-Bold", bold))
            return "GeorgiaLocal", "GeorgiaLocal-Bold"
    return "Times-Roman", "Times-Bold"


def para(text: str, style: ParagraphStyle) -> Paragraph:
    text = parse_inline(text)
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return Paragraph(text, style)


def build_pdf(md_path: Path, out_path: Path) -> None:
    blocks = parse_markdown(md_path)
    font, bold_font = register_pdf_fonts()
    base = getSampleStyleSheet()
    styles = {
        "body": ParagraphStyle(
            "Body",
            parent=base["BodyText"],
            fontName=font,
            fontSize=10.2,
            leading=13.0,
            spaceAfter=7,
            alignment=TA_LEFT,
        ),
        "h1": ParagraphStyle(
            "H1",
            parent=base["Heading1"],
            fontName=bold_font,
            fontSize=17,
            leading=21,
            textColor=colors.HexColor("#" + CHARCOAL),
            spaceBefore=15,
            spaceAfter=7,
        ),
        "h2": ParagraphStyle(
            "H2",
            parent=base["Heading2"],
            fontName=bold_font,
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#" + CRIMSON),
            spaceBefore=12,
            spaceAfter=5,
        ),
        "h3": ParagraphStyle(
            "H3",
            parent=base["Heading3"],
            fontName=bold_font,
            fontSize=11.5,
            leading=14,
            textColor=colors.HexColor("#" + CHARCOAL),
            spaceBefore=8,
            spaceAfter=4,
        ),
        "caption": ParagraphStyle(
            "Caption",
            parent=base["BodyText"],
            fontName=font,
            fontSize=8.5,
            leading=10,
            textColor=colors.HexColor("#" + MUTED),
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
        "cell": ParagraphStyle("Cell", parent=base["BodyText"], fontName=font, fontSize=6.2, leading=7.4),
        "header_cell": ParagraphStyle(
            "HeaderCell", parent=base["BodyText"], fontName=bold_font, fontSize=6.5, leading=7.8
        ),
    }

    story = []
    story.extend(
        [
            Spacer(1, 1.55 * inch),
            Paragraph(
                REPORT_TITLE.replace("&", "&amp;"),
                ParagraphStyle(
                    "CoverTitle",
                    parent=base["Title"],
                    fontName=bold_font,
                    fontSize=25,
                    leading=30,
                    alignment=TA_CENTER,
                    textColor=colors.HexColor("#" + CHARCOAL),
                    spaceAfter=10,
                ),
            ),
            Paragraph(
                REPORT_SUBTITLE.replace("&", "&amp;"),
                ParagraphStyle(
                    "CoverSubtitle",
                    parent=base["BodyText"],
                    fontName=font,
                    fontSize=13,
                    leading=16,
                    alignment=TA_CENTER,
                    textColor=colors.HexColor("#" + MUTED),
                    spaceAfter=18,
                ),
            ),
            Paragraph(
                REPORT_LABEL.replace("&", "&amp;"),
                ParagraphStyle(
                    "CoverLabel",
                    parent=base["BodyText"],
                    fontName=bold_font,
                    fontSize=10,
                    leading=12,
                    alignment=TA_CENTER,
                    textColor=colors.HexColor("#" + CRIMSON),
                    spaceAfter=26,
                ),
            ),
            Paragraph(
                REPORT_DATE,
                ParagraphStyle(
                    "CoverDate",
                    parent=base["BodyText"],
                    fontName=font,
                    fontSize=10,
                    leading=12,
                    alignment=TA_CENTER,
                    textColor=colors.HexColor("#" + MUTED),
                ),
            ),
            PageBreak(),
        ]
    )

    skip_title_block = True
    for block in blocks:
        if skip_title_block:
            if block.kind == "heading" and block.level == 1 and block.text == REPORT_TITLE:
                continue
            if block.kind == "paragraph" and (
                block.text.startswith("Subtitle:")
                or block.text.startswith("Header/report label:")
                or block.text.startswith("**Subtitle:**")
                or block.text.startswith("**Header/report label:**")
            ):
                continue
            skip_title_block = False

        if block.kind == "heading":
            if block.text == "Appendix":
                story.append(PageBreak())
            style = styles["h1"] if block.level == 1 else styles["h2"] if block.level == 2 else styles["h3"]
            story.append(para(block.text, style))
        elif block.kind == "paragraph":
            story.append(para(block.text, styles["body"]))
        elif block.kind == "list":
            bullet_type = "1" if block.ordered else "bullet"
            items = [ListItem(para(item, styles["body"]), leftIndent=14) for item in block.items or []]
            story.append(ListFlowable(items, bulletType=bullet_type, leftIndent=18, bulletFontName=font))
            story.append(Spacer(1, 3))
        elif block.kind == "image":
            img_path = md_path.parent / block.path
            with Image.open(img_path) as img:
                iw, ih = img.size
            max_w = 6.65 * inch
            max_h = 4.6 * inch
            scale = min(max_w / iw, max_h / ih)
            story.append(KeepTogether([RLImage(str(img_path), width=iw * scale, height=ih * scale), para(block.alt, styles["caption"])]))
        elif block.kind == "table":
            rows = block.rows or []
            if rows:
                cols = max(len(r) for r in rows)
                table_data = []
                for ri, row in enumerate(rows):
                    style = styles["header_cell"] if ri == 0 else styles["cell"]
                    table_data.append([para(row[ci] if ci < len(row) else "", style) for ci in range(cols)])
                col_width = 6.65 * inch / cols
                tbl = Table(table_data, colWidths=[col_width] * cols, repeatRows=1, splitByRow=1)
                tbl.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F2F2")),
                            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CCCCCC")),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 3),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                            ("TOPPADDING", (0, 0), (-1, -1), 3),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                        ]
                    )
                )
                story.append(tbl)
                story.append(Spacer(1, 9))
        elif block.kind == "rule":
            story.append(Spacer(1, 5))

    def page_canvas(canvas, doc):
        canvas.saveState()
        canvas.setFont(font, 8.5)
        canvas.setFillColor(colors.HexColor("#" + MUTED))
        canvas.drawString(doc.leftMargin, 0.43 * inch, REPORT_LABEL)
        canvas.drawRightString(letter[0] - doc.rightMargin, 0.43 * inch, f"Page {doc.page}")
        canvas.restoreState()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=letter,
        rightMargin=0.82 * inch,
        leftMargin=0.82 * inch,
        topMargin=0.82 * inch,
        bottomMargin=0.75 * inch,
        title=REPORT_TITLE,
        author="gabriel-wages project",
    )
    doc.build(story, onFirstPage=page_canvas, onLaterPages=page_canvas)


def export(md_path: Path, output_dir: Path) -> tuple[Path, Path, list[Path]]:
    if not md_path.exists():
        raise FileNotFoundError(md_path)
    images = check_images(md_path)
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / "deeper_look_safety_non_safety_wage_mechanisms_2026-07-10.docx"
    pdf_path = output_dir / "deeper_look_safety_non_safety_wage_mechanisms_2026-07-10.pdf"
    build_docx(md_path, docx_path)
    build_pdf(md_path, pdf_path)
    return docx_path, pdf_path, images


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    sub = parser.add_subparsers(dest="command", required=True)

    p_int = sub.add_parser("integrate", help="Combine reviewed scaffold and appendix into one Markdown report")
    p_int.add_argument("--scaffold", required=True, type=Path)
    p_int.add_argument("--appendix", required=True, type=Path)
    p_int.add_argument("--output", required=True, type=Path)

    p_exp = sub.add_parser("export", help="Export integrated Markdown to DOCX and PDF")
    p_exp.add_argument("--input", required=True, type=Path)
    p_exp.add_argument("--output-dir", required=True, type=Path)

    p_check = sub.add_parser("check-tools", help="Report available local export tools")

    args = parser.parse_args(argv)
    if args.command == "check-tools":
        print("pandoc:", shutil.which("pandoc") or "missing")
        print("soffice:", shutil.which("soffice") or "missing")
        print("python-docx: available")
        print("reportlab: available")
        return 0
    if args.command == "integrate":
        integrate_markdown(args.scaffold, args.appendix, args.output)
        print(args.output)
        return 0
    if args.command == "export":
        docx_path, pdf_path, images = export(args.input, args.output_dir)
        print(f"DOCX: {docx_path} ({docx_path.stat().st_size} bytes)")
        print(f"PDF: {pdf_path} ({pdf_path.stat().st_size} bytes)")
        print(f"Images: {len(images)}")
        return 0
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
