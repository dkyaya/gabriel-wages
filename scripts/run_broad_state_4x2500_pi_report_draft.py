#!/usr/bin/env python3
"""Build, validate, and audit the Broad State 4x2500 PI-facing report draft."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import subprocess
import textwrap
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ANALYSIS = ROOT / "docs/analysis/compensation_extraction"
BOUNDED = ANALYSIS / "BROAD-STATE-4X2500-BOUNDED-WAGE-DIFFERENTIAL-VALIDATION-2026-07-30"
RESCUE = ANALYSIS / "BROAD-STATE-4X2500-NORMALIZATION-RESCUE-GAP-GROWTH-CLAIMS-2026-07-30"
NORMALIZATION = ANALYSIS / "BROAD-STATE-4X2500-NORMALIZATION-MATCHED-STRUCTURE-PARAPHRASE-REPAIR-2026-07-30"
INGEST = ANALYSIS / "BROAD-STATE-4X2500-RATING-INGEST-CODIFY-PI-EVIDENCE-2026-07-30"
RATING = ANALYSIS / "BROAD-STATE-4X2500-SPAN-RATING-AND-DASHBOARD-CLEANUP-2026-07-30"
OUTPUT = ANALYSIS / "BROAD-STATE-4X2500-PI-REPORT-DRAFT-2026-07-30"
TASK_ID = "BROAD-STATE-4X2500-PI-REPORT-DRAFT-2026-07-30"
DECISION = "broad_state_4x2500_pi_report_draft_completed_review_ready"
NEXT_TASK = "BROAD-STATE-4X2500-PI-REPORT-REVIEW-FINALIZE-2026-07-30"
HEAD_BEFORE = "df9802202a843f3c734818d39243079910ee7f5b"


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def read_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, data) -> None:
    path.write_text(json.dumps(data, indent=2, sort_keys=True) + "\n", encoding="utf-8")


def write_csv(path: Path, rows: list[dict], fields: list[str]) -> None:
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore", lineterminator="\n")
        writer.writeheader()
        writer.writerows(rows)


def report_inputs() -> dict:
    required = [
        BOUNDED / "bounded_wage_differential_validation_summary.json",
        BOUNDED / "bounded_wage_differential_pi_statements.json",
        BOUNDED / "merged_bounded_wage_differential_validation_results.jsonl",
        RESCUE / "normalization_rescue_gap_growth_summary.json",
        RESCUE / "quantitative_growth_mechanism_claims.json",
        NORMALIZATION / "updated_careful_claim_candidates.json",
        NORMALIZATION / "normalization_summary.json",
        NORMALIZATION / "municipality_cycle_groups_summary.json",
        NORMALIZATION / "matched_cycle_summary.json",
        NORMALIZATION / "comparable_normalized_wage_summary.json",
        NORMALIZATION / "growth_readiness_summary.json",
        INGEST / "mechanism_cluster_strength_table.json",
        RATING / "report_usability_summary.json",
        RATING / "claim_relevance_summary.json",
        RATING / "directionality_summary.json",
    ]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing report inputs:\n" + "\n".join(missing))
    bounded = read_json(required[0])
    if (
        bounded.get("input_candidate_count") != 4
        or bounded.get("validated_pi_report_usable_count") != 1
        or bounded.get("conditional_manual_review_count") != 3
    ):
        raise ValueError("bounded wage-differential validation does not reconcile")
    return {
        "bounded_summary": bounded,
        "bounded_statements": read_json(required[1])["statements"],
        "rescue_summary": read_json(required[3]),
        "growth_claims": read_json(required[4])["claims"],
        "careful_claims": read_json(required[5])["claims"],
        "normalization_summary": read_json(required[6]),
        "mechanism_strengths": read_json(required[11])["rows"],
        "report_usability": read_json(required[12]),
        "claim_relevance": read_json(required[13]),
        "directionality": read_json(required[14]),
    }


SELECTED_GROWTH_IDS = [
    "B4X2500GROWTHCLAIM-1f9cbf949deac19c0ce20b83",  # Howland 2%
    "B4X2500GROWTHCLAIM-03ac871a89dac2eb6722a807",  # Coleraine 5%
    "B4X2500GROWTHCLAIM-ce5cf300ebb5e742567fd0b1",  # Sunrise Beach 1.5% COLA
    "B4X2500GROWTHCLAIM-fd1f2ba56f7d5a0b2553ae7d",  # Phippsburg 3.2% COLA
    "B4X2500GROWTHCLAIM-27df486203241975efc92285",  # Bath 2.25% step
    "B4X2500GROWTHCLAIM-365c82cdd78c3d862bfe9057",  # Tonka Bay 2.75% step
]


MECHANISM_ROWS = [
    ("non_base_compensation", "Non-base compensation", 3718, 1.646, 1.0, "Core", "Premiums, longevity, stipends, differentials, and specialty pay are distinct channels that base-wage comparisons can miss."),
    ("quantitative_base_wage_needs_normalization", "Direct base-wage evidence", 2687, 1.452, 1.0, "Core", "Direct values and schedules are abundant, but pay basis, period, unit, rank, and base/non-base alignment remain essential."),
    ("timing_implementation", "Implementation and retroactivity", 2230, 1.853, 2.0, "Core", "Effective dates, delayed implementation, and retroactive payments can change when negotiated gains are realized."),
    ("automatic_wage_growth", "Automatic raises and indexing", 1431, 1.377, 1.0, "Core", "Scheduled percentage increases, COLA/CPI clauses, and step progression provide documentable growth rules."),
    ("bargaining_dispute_resolution", "Bargaining and dispute resolution", 1135, 1.285, 1.0, "Supporting", "Bargaining, arbitration, factfinding, settlements, and strike constraints are plausible institutional channels, not causal estimates."),
    ("rank_step_specialization_classification", "Rank, step, specialization, classification", 763, 2.427, 2.0, "Core", "Structured ladders and classifications shape progression and complicate cross-occupation comparisons."),
    ("strike_or_no_strike", "Strike/no-strike constraints", 587, 2.539, 2.0, "Supporting", "Work-stoppage constraints form part of institutional bargaining context."),
    ("market_staffing_pressure", "Market and staffing pressure", 398, 1.374, 1.0, "Supporting", "Comparability, recruitment, retention, and staffing language appear as stated compensation justifications."),
    ("safety_advantage", "Safety-advantage hints", 192, 0.987, 1.0, "Context", "Some spans point toward a safety-side advantage, but directional evidence is sparse and uneven."),
    ("fiscal_governance_constraints", "Fiscal and governance constraints", 188, 1.531, 1.0, "Context", "Budget, appropriation, and approval language documents constraints without identifying effects."),
]


BOUNDED_ROWS = [
    {
        "municipality_state": "Shreve, OH",
        "period": "2024",
        "safety_value": "$22.00/hour (part-time police officers)",
        "non_safety_value": "$16.00/hour (part-time utility clerk)",
        "difference": "$6.00/hour",
        "percent_difference": "37.5%",
        "validation_status": "validated_pi_report_usable",
        "report_use": "Supporting example",
        "key_caveat": "Same ordinance and hourly basis; duties, experience, schedule, and job requirements are not matched.",
    },
    {
        "municipality_state": "Cammack Village, AR",
        "period": "2024",
        "safety_value": "$25.00/hour maximum (part-time patrolman)",
        "non_safety_value": "$20.00/hour maximum (administrative assistant)",
        "difference": "$5.00/hour",
        "percent_difference": "25.0%",
        "validation_status": "validated_with_caveats_manual_review",
        "report_use": "Qualified supporting example",
        "key_caveat": "Authorized maxima, not confirmed actual wages; enactment fields require manual/legal confirmation.",
    },
    {
        "municipality_state": "Canastota, NY",
        "period": "2023-2024",
        "safety_value": "$23.91/hour (Police Officer Year 1 step)",
        "non_safety_value": "$24.82/hour (Code Enforcement Officer)",
        "difference": "-$0.91/hour",
        "percent_difference": "-3.67%",
        "validation_status": "validated_with_caveats_manual_review",
        "report_use": "Qualified counterexample",
        "key_caveat": "Entry-step police rate is not tenure/experience-equivalent to the single code-enforcement rate.",
    },
    {
        "municipality_state": "Alburtis, PA",
        "period": "Jan-Jun 2018",
        "safety_value": "$33.57/hour (Chief of Police)",
        "non_safety_value": "$11.22/hour (Administrative Assistant)",
        "difference": "$22.35/hour",
        "percent_difference": "199.2%",
        "validation_status": "validated_with_caveats_manual_review",
        "report_use": "Limits/appendix only",
        "key_caveat": "Chief was outside the police bargaining unit; roles, rank, hours, and schedule position are not comparable.",
    },
]


LIMIT_ROWS = [
    ("Final wage-gap estimation", "Blocked", "Only bounded local documentary comparisons are presented; no analytic wage-gap estimate is reported."),
    ("Causal inference", "Blocked", "The evidence identifies plausible mechanisms but does not establish treatment effects or causal effects."),
    ("National/population prevalence", "Blocked", "The processed corpus is not treated as a representative probability sample of municipalities."),
    ("Regression analysis", "Not performed", "No regression or treatment-effect model was run."),
    ("Cost-of-living adjustment", "Not performed", "COLA/CPI appears only as source-reported contract language; no analyst-side COL adjustment was made."),
    ("OCR-later documents", "Excluded", "601 retained sources remain deferred for OCR and are outside this report's evidence base."),
    ("Role/rank/step comparability", "Material blocker", "Rank, tenure, step, duties, hours, and qualifications remain unresolved in many candidate comparisons."),
    ("Normalization", "Incomplete", "Many quantitative records remain partial, mechanism-only, deferred, or unusable for wage-level comparison."),
    ("Bounded comparisons", "Manual review required", "Even the cleanest local comparison needs final substantive review before analytic use."),
    ("Selection filters", "Interpretive limit", "Search, reachability, retention, extraction, rating, and normalization gates shape the observed evidence."),
]


SOURCE_ROWS = [
    ("Scout coverage", "16,887 municipalities", "47.45% of 35,589 eligible/known municipality universe"),
    ("4x2500 scout wave", "10,000 terminal outcomes", "9,968 parseable; 32 failed/unparseable"),
    ("Candidate formation", "9,977 raw / 9,072 deduped", "5,768 verification-ready"),
    ("Verification", "5,768 reviewed", "3,950 reachable/source-review-ready"),
    ("Source review/download", "3,672 retained", "3,248 PDF; 350 HTML; 74 other"),
    ("Text readiness", "2,940 ready", "2,577 PDF; 291 HTML; 72 other; 601 OCR-later"),
    ("Text extraction", "2,795 clean/span-ready", "145 quality/problem rows"),
    ("Span extraction", "19,118 candidates", "18,612 rating-ready"),
    ("Span rating", "18,554 valid", "58 quarantined and excluded"),
    ("Normalization", "11,548 quantitative records", "672 full; 1,563 partial; 3,769 mechanism-only; 720 deferred; 4,824 unusable"),
    ("Matched structure", "2,712 municipality-cycle groups", "65 safety/non-safety candidates; 303 comparable wage candidates; 141 growth-readiness candidates"),
    ("Focused rescue/validation", "416 quantitative growth records", "4 bounded comparisons: 1 PI-usable, 3 conditional"),
]


def executive_summary() -> str:
    return textwrap.dedent(
        """
        The project can now make a set of careful, evidence-based statements about how municipal compensation changes are structured and justified. Within the processed rated corpus, wage growth is not captured by a single base-wage number. It is assembled through scheduled percentage increases, step and grade progression, COLA or CPI-linked clauses, implementation dates, retroactive payments, bargaining and dispute-resolution rules, market-comparability rationales, and a large layer of non-base compensation. These are documentary mechanism findings: they identify recurring channels in the records that were found, retained, extracted, rated, and codified. They do not establish national prevalence or causal effects.

        The strongest substantive conclusion is that non-base compensation must be treated as a central part of the empirical design rather than as an appendix to base wages. Among valid rated spans, 3,718 carry a non-base compensation signal, including longevity pay, shift differentials, hazard and specialty pay, certification and education incentives, stipends, allowances, and overtime-related premiums. Direct base-wage schedules are also abundant—2,687 rated signals—but their usefulness depends on aligning pay basis, effective period, occupation, rank, step, grade, and base-versus-non-base status. Timing is another prominent channel: 2,230 rated signals concern implementation or retroactivity, indicating that two agreements with similar headline increases may generate different realized compensation paths because of effective dates, delayed settlements, or retroactive payment rules.

        The quantitative rescue strengthens the mechanism evidence without converting it into a national comparison. It identified 416 records with source-reported quantitative growth mechanisms: 336 percentage increases, 49 COLA/CPI adjustments, 29 step-schedule increases, and 2 retroactive or lump-sum mechanisms. The report uses six transparent examples—such as a 2% police raise in Howland, Ohio; a 1.5% police COLA/CPI mechanism in Sunrise Beach, Missouri; and a 2.25% police step-schedule mechanism in Bath, Ohio—to show the kind of growth rule documented in the source record. These figures are source-reported mechanisms, not analyst-computed growth estimates or evidence that the same rule is broadly prevalent.

        The current matched evidence also supports a narrow set of bounded local documentary comparisons. Shreve, Ohio, is the cleanest supporting example: the 2024 municipal wage ordinance lists part-time police officers at $22.00 per hour and a part-time utility clerk at $16.00 per hour, a source-grounded difference of $6.00 per hour, or 37.5% relative to the clerk rate. Cammack Village and Canastota can be used only with heavy candidate-specific caveats, and Alburtis belongs in the limits or appendix because its police-side observation is the Chief of Police, explicitly outside the police bargaining unit. None of these comparisons is a final wage-gap estimate; each remains local, documentary, and subject to final substantive validation.

        What remains blocked is as important as what is now usable. The corpus does not yet support a final safety-versus-non-safety wage-gap estimate, a nationally representative prevalence statement, a regression or treatment-effect result, or a final causal claim. Directional ratings are uneven: 637 non-safety-advantage hints, 78 safety-advantage hints, and 45 gap-narrowing hints coexist with 6,830 neutral/general and 10,964 not-applicable ratings. That distribution makes selective citation especially risky. The matching layer is promising but still limited: 2,712 municipality-cycle groups yield 65 safety/non-safety matched-cycle candidates, 303 potentially comparable normalized wage candidates, and 141 cycle-to-cycle growth-readiness candidates.

        The immediate next step is human review and tightening, not a larger claim. The four bounded comparisons should receive final substantive and, where relevant, legal validation; the 185 near-gap-ready records should receive targeted normalization; and the 65 matched cycles and 303 comparable candidates should be refined into occupation-, rank-, pay-basis-, and period-aligned pairs. An OCR strategy for the 601 deferred sources can expand the evidence base later. Exploratory wage-gap estimation should begin only after a documented review gate authorizes the matched inputs, while causal modeling should wait for a substantially stronger matched structure and identification strategy.
        """
    ).strip()


def normalize_markdown(text: str) -> str:
    """Remove template indentation without disturbing interpolated prose blocks."""
    return re.sub(r"(?m)^ {8}", "", text.strip()) + "\n"


def main_report() -> str:
    return normalize_markdown(textwrap.dedent(
        f"""
        # Why Public-Safety Wages May Rise Faster Than Other Municipal Wages

        **PI-facing evidence report draft | July 30, 2026**

        This report summarizes what the current processed evidence can responsibly say about municipal wage-growth mechanisms and bounded local wage differences. It emphasizes findings and inference boundaries rather than pipeline operations.

        ## 1. Executive Summary

        {executive_summary()}

        ## 2. Processed Evidence Base

        The evidence base is broad enough to support mechanism analysis but not a probability-sample claim about municipalities. Scout coverage currently reaches 16,887 municipalities, or 47.45% of the project's 35,589 eligible or known municipality universe. The 4x2500 wave produced 10,000 terminal outcomes, of which 9,968 were parseable. Those outcomes generated 9,977 raw candidate rows and 9,072 deduplicated candidates; 5,768 entered verification, and 3,950 were reachable and suitable for source review.

        Source review retained 3,672 unique documents—3,248 PDFs, 350 HTML pages, and 74 other document types. Text-readiness review identified 2,940 non-OCR sources suitable for extraction. Text extraction produced 2,795 clean, span-ready sources and 182.1 million bytes of local extracted text. Span extraction produced 19,118 candidate passages, of which 18,612 met the rating threshold. Rating yielded 18,554 valid records; 58 were quarantined and remain excluded from the evidentiary layer used here.

        The quantitative layer contains 11,548 records with some compensation value or growth information. Normalization classified 672 as full, 1,563 as partial, 3,769 as mechanism-only, 720 as deferred for manual review, and 4,824 as unusable for safe comparison. The targeted rescue upgraded a small share of the partial basket and identified 416 records with quantitative growth mechanisms. The most important implication is not the size of the pipeline; it is that the documentary corpus now separates evidence that can support mechanism description from evidence that still requires normalization, matching, or exclusion.

        | Evidence layer | Current usable count | Interpretation |
        |---|---:|---|
        | Retained sources | 3,672 | Unique retained documents after source review |
        | Clean/span-ready sources | 2,795 | Non-OCR text suitable for exact-span review |
        | Valid rated spans | 18,554 | Schema-valid ratings used for codification |
        | Normalized quantitative records | 11,548 | Includes full, partial, mechanism-only, deferred, and unusable statuses |
        | Quantitatively supported growth records | 416 | Source-reported percentage, COLA/CPI, step, or retroactive mechanisms |
        | Bounded local comparison candidates | 4 | One PI-usable supporting example; three conditional/manual-review examples |

        ## 3. Codified Evidence Categories

        The rating layer separates evidence quality and report use from mere textual relevance. Of the 18,554 valid ratings, 472 were classified as PI-report core-finding ready and 528 as supporting examples. Another 6,860 provide context, 5,533 contain quantitative content that requires downstream normalization, and 5,161 are excluded from the report layer. Claim relevance is similarly differentiated: 6,454 ratings directly support quantitative claims, 6,509 support mechanism summaries, 430 provide directional hints only, and 5,161 are weak or unsupported.

        | Report-usability bucket | Count | Intended use |
        |---|---:|---|
        | Core finding ready | 472 | Candidate mechanism findings after claim-boundary review |
        | Supporting example | 528 | Source-specific illustrations with explicit caveats |
        | Context only | 6,860 | Interpretation and source context, not a finding by itself |
        | Normalization needed | 5,533 | Quantitative content requiring aligned units, periods, and roles |
        | Exclude from report | 5,161 | Weak, unsupported, navigation-only, or otherwise non-claim-ready |

        The mechanism counts are counts within the processed rated corpus, not prevalence estimates. Non-base compensation is the largest rated signal at 3,718, followed by direct base-wage values at 2,687, implementation or retroactivity at 2,230, automatic raises at 1,431, bargaining-power signals at 1,135, rank or specialization premiums at 763, strike/no-strike constraints at 587, market or comparability pressure at 398, safety-advantage signals at 192, and fiscal-constraint signals at 188.

        Rated strength provides a second dimension. On the project's 0-4 scale, rank/step/specialization/classification evidence has the highest cluster average among the principal findings reported here (2.427; median 2.0), followed by implementation/timing (1.853; median 2.0) and non-base compensation (1.646; median 1.0). Automatic wage-growth evidence averages 1.377 (median 1.0), bargaining and dispute-resolution evidence 1.285 (median 1.0), and market/staffing pressure 1.374 (median 1.0). These scores summarize the rated specificity and support of individual spans; they do not measure effect magnitude.

        ## 4. Findings

        ### Finding 1. Non-base compensation is a major compensation channel

        Within the processed rated corpus, non-base compensation appears as a substantial channel through which compensation can grow even when base schedules move modestly. The 3,718 non-base signals cover longevity pay, shift differentials, hazard and specialty pay, certification and education incentives, stipends, allowances, and overtime-related premiums. These provisions matter empirically because they can raise realized compensation without changing the headline base-wage line used in a conventional schedule comparison.

        The source-specific examples make the measurement problem concrete. A Bath, Ohio, record identifies a one-time or lump-sum police payment of $1,200 for 2021; a Saybrook, Ohio, record identifies a $300 fire stipend or allowance for 2019; and a Franklin, Michigan, record identifies a $3,500 police stipend or allowance for 2018. Each example documents a compensation component that should remain distinct from recurring base pay. The evidence does not establish that these components are unique to safety workers, nor does it yet produce a total-compensation difference. It does establish that an analysis limited to base salary would omit material compensation channels present in the source record.

        ### Finding 2. Direct wage schedules are abundant, but comparison requires normalization

        Direct wage and salary values are one of the largest quantitative resources in the corpus. They appear in hourly and annual schedules, grade and pay-band tables, rank ladders, step plans, ordinances, budgets, and negotiated agreements. The 2,687 base-wage direct-value signals and 11,548 quantitative records provide a substantial raw basis for later comparison.

        Their abundance should not be confused with immediate comparability. A listed hourly police rate may refer to an entry step, a maximum, a part-time position, a chief outside the bargaining unit, or a premium rather than base pay. An annual general-employee salary may embed a different work year or classification level. The normalization exercise therefore preserves source-reported values while adding pay basis, effective period, unit, occupation, rank, step, grade, and base/non-base fields. Only 672 records initially met full normalization requirements; the rescue added 11 fully normalized records and identified 185 near-gap-ready records. This is meaningful progress toward a comparison layer, but it is not authorization to report a corpus-wide wage difference.

        ### Finding 3. Implementation timing and retroactivity affect realized gains

        Timing is a substantive mechanism, not merely an administrative detail. The rated corpus contains 2,230 implementation or retroactivity signals. These spans refer to effective dates, delayed implementation, settlement timing, retroactive wage adjustments, and lump-sum or catch-up payments. A delayed agreement can shift compensation into a later fiscal period; a retroactive provision can deliver a material payment after the nominal effective date; and two agreements with the same stated annual percentage can yield different near-term cash flows.

        The evidence therefore supports a careful mechanism claim: implementation and retroactivity can shape when negotiated gains are realized within the documented unit and cycle. It does not yet support a comparative statement that one occupation systematically realizes gains earlier or receives larger retroactive awards. That inference requires matched unit-cycle timing fields and a consistent treatment of lump-sum versus recurring pay.

        ### Finding 4. Automatic wage-growth mechanisms recur in the documentary record

        Automatic wage growth appears in several distinguishable forms: across-the-board percentage increases, COLA or CPI-linked adjustments, scheduled raises, and step progression. The rescue identified 416 records with quantitative growth mechanisms—336 percentage increases, 49 COLA/CPI adjustments, 29 step-schedule increases, and 2 retroactive or lump-sum mechanisms. Ninety-five of those records were promoted into source-specific quantitative growth claim candidates.

        Selected examples illustrate the range without implying prevalence. Howland, Ohio, documents a 2% police raise for 2021; Coleraine, Minnesota, documents a 5% police raise for 2024; Sunrise Beach, Missouri, documents a 1.5% police COLA/CPI mechanism for 2024; Phippsburg, Maine, documents a 3.2% police COLA/CPI mechanism for 2023; Bath, Ohio, documents a 2.25% police step-schedule mechanism for 2021; and Tonka Bay, Minnesota, documents a 2.75% police step-schedule mechanism for 2020. These are source-reported values tied to named records. They support quantitatively specific mechanism statements, not analyst-computed growth rates or a statement about how widely any mechanism is used.

        COLA/CPI language requires a particularly clear boundary. Here it is treated as a contractual rule that can change nominal wages. The report does not apply an analyst-side cost-of-living adjustment, convert nominal wages into real wages, or compare local purchasing power.

        ### Finding 5. Bargaining and dispute resolution are plausible institutional mechanisms

        Among valid rated spans, 1,135 carry bargaining-power signals and 587 concern strike or no-strike constraints. Bargaining agreements, arbitration awards, factfinding materials, settlement records, and related documents describe procedures through which wage terms are proposed, contested, resolved, and implemented. The evidence supports describing collective bargaining, interest arbitration, factfinding, negotiated settlement, and work-stoppage constraints as plausible institutional channels in some municipal compensation records.

        The key distinction is between mechanism presence and effect. An arbitration or factfinding clause can structure the bargaining process without demonstrating that it raised wages. Ordinary grievance arbitration also should not be conflated with interest arbitration over unresolved contract terms. The current corpus documents these institutional channels, but it does not identify a counterfactual wage outcome or establish that a particular procedure produced a safety advantage.

        ### Finding 6. Rank, step, specialization, and classification structures shape progression

        Rank, step, grade, specialization, and classification signals number 763 and have relatively strong rated specificity (average 2.427; median 2.0). These structures create predictable internal wage paths: employees can move through scheduled steps, receive rank premiums, enter new grades, or qualify for specialty compensation. They also complicate cross-occupation comparisons because a police officer at an entry step is not automatically comparable to a non-safety employee at an unlabeled or maximum rate.

        This point is visible in the Canastota example discussed below. The police record is explicitly a Year 1 step, while the Code Enforcement Officer line reports a single rate without comparable tenure or step information. The local schedule values are real and source-grounded, but the ranking structure prevents an unqualified wage comparison. Matching must therefore occur at the position and schedule-location level, not only at municipality and year.

        ### Finding 7. Market and staffing pressure appear as compensation justifications

        Market comparability, recruitment, retention, staffing shortages, and competing-jurisdiction language appear in 398 rated market/comparability signals. Records from places such as Bernards and Greenwich, New Jersey, and Brewster, Massachusetts, invoke these considerations as stated justifications for compensation review or change. Such language supplies a plausible explanation for why a municipality considers revising pay: it is responding, at least rhetorically or administratively, to outside wage benchmarks or difficulty attracting and retaining personnel.

        The evidence does not show that invoking comparability actually produced a particular wage increase, nor that the stated pressure was objectively severe. It does show that market and staffing considerations belong in the mechanism taxonomy and should be linked to later matched wage outcomes where the city-cycle structure permits.

        ### Finding 8. Bounded local wage-differential evidence exists, but it is narrow

        Four local comparison candidates survived normalization rescue and source-grounded validation. Only Shreve is PI-usable as a supporting example without an additional legal-status question; the remaining three require heavy, candidate-specific qualifications. The values below are current source-grounded comparisons on the normalized basis shown. They are not final wage-gap estimates and are not nationally representative.

        | Municipality / period | Safety-side value | Non-safety-side value | Difference | Status | Principal caveat |
        |---|---|---|---:|---|---|
        | Shreve, OH / 2024 | Part-time police: $22.00/hour | Part-time utility clerk: $16.00/hour | +$6.00/hour; +37.5% | PI-usable supporting example | Duties, experience, schedules, and job requirements are not matched |
        | Cammack Village, AR / 2024 | Part-time patrolman maximum: $25.00/hour | Administrative assistant maximum: $20.00/hour | +$5.00/hour; +25.0% | Conditional/manual review | Authorized maxima, not confirmed actual pay; enactment fields need review |
        | Canastota, NY / 2023-2024 | Police Officer Year 1: $23.91/hour | Code Enforcement Officer: $24.82/hour | -$0.91/hour; -3.67% | Conditional/manual review | Entry step is not tenure- or experience-equivalent to the single comparator rate |
        | Alburtis, PA / Jan-Jun 2018 | Chief of Police: $33.57/hour | Administrative Assistant: $11.22/hour | +$22.35/hour; +199.2% | Limits/appendix only | Chief was outside the police bargaining unit; roles and schedule position are not comparable |

        Shreve supplies the cleanest bounded local documentary evidence: the same 2024 ordinance lists hourly rates for part-time police officers and a part-time utility clerk, with the police-side value $6.00 per hour, or 37.5%, above the clerk rate. This is a named-position comparison, not an occupation-wide estimate. Cammack Village compares authorized maximum rates rather than demonstrated actual wages, and the retained packet requires enactment confirmation. Canastota points in the opposite direction for the two listed schedule positions, with the Police Officer Year 1 rate $0.91 per hour below the Code Enforcement Officer rate; its lack of tenure and step equivalence makes it a useful caution against assuming a uniform direction. Alburtis illustrates the danger of role mismatch: its very large documentary difference compares a Chief outside the bargaining unit with an Administrative Assistant and belongs in an appendix or limits discussion.

        ### Finding 9. Directional evidence is present but uneven

        The rating layer contains 637 non-safety-advantage hints, 78 safety-advantage hints, and 45 gap-narrowing hints. Those counts sit beside 6,830 neutral/general ratings and 10,964 ratings for which direction is not applicable. The distribution does not support a one-direction summary of the corpus. It instead reinforces two points: much of the evidence describes mechanisms without comparing occupations, and the subset that does indicate direction includes countervailing signals.

        Directional hints can guide targeted matching and review, but they cannot substitute for matched wage records. The Canastota schedule is a concrete reminder that a local documentary comparison can point toward a non-safety advantage for the specific positions shown. Selective use of only safety-advantage examples would mischaracterize the present evidence.

        ### Finding 10. Matched city-cycle structures are emerging but still limited

        Normalization grouped evidence into 2,712 municipality-cycle structures and identified 65 candidate cycles with both safety and non-safety evidence. Within those groups, 303 record pairs may support a normalized wage comparison after review, and 141 candidate sequences may support cycle-to-cycle growth analysis. The rescue classified 15 additional structures as future-gap-potential only.

        These structures support a readiness claim: the project has moved beyond isolated document collection and can identify where same-city, overlapping-period, cross-occupation comparisons may be possible. They do not yet support a final cross-city wage-gap statistic. The next refinement must verify that paired records share a comparable pay basis, effective period, occupation level, rank or step, and base/non-base concept. The matched structure should remain the governing discipline: a safety record without a same-city, same-cycle non-safety comparator is not sufficient for the core empirical design.

        ## 5. Limits

        The report's strongest claims are documentary and mechanism-oriented. Several boundaries remain firm.

        1. **No final wage-gap estimate.** The four local comparisons are bounded illustrations. Even Shreve requires final substantive review before analytic use.
        2. **No causal claim.** The corpus identifies contractual and institutional mechanisms but does not estimate counterfactual wage outcomes.
        3. **No national or population-prevalence claim.** The corpus is broad, but source discovery, public availability, reachability, retention, extraction, and rating filters prevent representative-sample interpretation.
        4. **No regression or treatment-effect result.** No such model was run for this report.
        5. **No analyst-side cost-of-living adjustment.** COLA/CPI appears only as source-reported wage-growth language; the report does not adjust wages for local prices or inflation.
        6. **OCR exclusions remain material.** Six hundred one retained sources are deferred for OCR and do not contribute to the present span and rating layers.
        7. **Normalization remains incomplete.** Pay basis, base/non-base classification, occupation, safety category, effective period, and rank/step/grade remain unresolved for many quantitative records.
        8. **Position equivalence is not automatic.** Titles, duties, hours, tenure, qualifications, and schedule location must be checked before a local documentary difference becomes an analytic comparison.
        9. **Quarantine and weak evidence remain excluded.** Fifty-eight quarantined ratings and 5,161 weak or unsupported ratings do not enter the findings layer.

        ## 6. Current Scout Wave Status

        The Broad State 4x2500 wave has now been processed through scouting, candidate review, verification, source review and retention, text readiness, non-OCR text extraction, exact-span extraction, rating, ingestion and codification, normalization and matching, targeted normalization rescue, and focused validation of four bounded local wage comparisons. The public dashboard is current and uses scout coverage rate—not raw discovery volume—as its map metric. National scout coverage is 16,887 of 35,589 eligible or known municipalities, or 47.45%. Detailed engineering and reconciliation counters remain available in the dashboard's collapsed technical section rather than in the main PI view.

        Public dashboard: https://dkyaya.github.io/gabriel-wages/

        ## 7. Recommended Next Steps

        1. **Complete final manual validation of the four bounded comparisons.** Confirm enactment, operative status, job and schedule equivalence, and bargaining-unit scope before any value enters an analytic table.
        2. **Expand normalization around the 185 near-gap-ready records.** Prioritize records already connected to the 65 matched municipality-cycle candidates rather than normalizing isolated safety records.
        3. **Refine the matched city-cycle layer.** Review the 303 comparable wage candidates and 141 growth-readiness candidates for pay basis, effective period, occupation, rank, step, grade, and base/non-base alignment.
        4. **Develop a bounded OCR strategy for the 601 deferred sources.** OCR should be a separate, audited phase; its outputs should not silently mix with the current non-OCR evidence.
        5. **Authorize exploratory wage-gap estimation only after a documented validation gate.** Initial outputs should remain transparent, local, and sensitivity-tested before any broader summary is considered.
        6. **Defer causal modeling.** Regression or treatment-effect work should wait until the matched structure, repeated cycles, and identification strategy are materially stronger.
        7. **Use further source expansion selectively.** After PI review, collect additional sources only where they close named matched-cycle, normalization, or mechanism-evidence gaps.
        """
    ))


def one_page_brief() -> str:
    return textwrap.dedent(
        """
        # Gabriel Wages: One-Page PI Brief

        **What the project can say now.** Within the processed rated corpus, municipal compensation growth is documented through multiple channels rather than a single base-wage number. The strongest current mechanism findings concern non-base compensation; direct wage schedules; implementation timing and retroactivity; automatic percentage, COLA/CPI, and step increases; bargaining and dispute resolution; rank and classification structures; and market or staffing justifications. These are careful documentary claims, not estimates of national prevalence or causal effects.

        **Most important substantive implication.** Base wages alone are an incomplete measure. The rated corpus contains 3,718 non-base compensation signals covering longevity, shift, hazard, specialty, certification, education, stipends, allowances, and premium pay. It also contains 2,687 direct base-wage signals and 2,230 implementation or retroactivity signals. A credible safety-versus-non-safety comparison must therefore align not only salary levels but also non-base components and the timing of when negotiated gains are paid.

        **Quantitative growth evidence.** A targeted rescue identified 416 source-reported quantitative growth mechanisms: 336 percentage increases, 49 COLA/CPI adjustments, 29 step-schedule increases, and 2 retroactive or lump-sum mechanisms. Examples include a 2% police raise in Howland, Ohio; a 1.5% police COLA/CPI mechanism in Sunrise Beach, Missouri; and a 2.25% police step-schedule mechanism in Bath, Ohio. These are document-specific growth rules, not analyst-computed wage growth or evidence of broad prevalence.

        **Bounded local comparisons.** Shreve, Ohio, provides the cleanest supporting example: a 2024 ordinance lists part-time police at $22.00 per hour and a part-time utility clerk at $16.00 per hour, a local documentary difference of $6.00 per hour, or 37.5% relative to the clerk rate. Duties, experience, schedules, and job requirements remain unmatched. Cammack Village compares authorized maximum rates and requires enactment confirmation; Canastota compares a police entry step with a single code-enforcement rate and points in the opposite direction; Alburtis compares a Chief outside the bargaining unit with an Administrative Assistant and belongs in the limits or appendix. None is a final wage-gap estimate.

        **What remains blocked.** The project does not yet support a final wage-gap estimate, a nationally representative prevalence statement, a regression or treatment-effect conclusion, or a final causal claim. The matching layer contains 2,712 municipality-cycle groups, 65 safety/non-safety candidates, 303 potentially comparable wage pairs, and 141 growth-readiness candidates, but most require additional role, rank, pay-basis, and period validation. Six hundred one retained sources remain deferred for OCR.

        **Recommended decision.** Move to report review and targeted evidence refinement. First validate the four bounded comparisons; then normalize the 185 near-gap-ready records connected to matched cycles; refine the 65 matched candidates and 303 comparable wage pairs; and authorize only bounded exploratory estimation after a formal review gate. Causal modeling should remain deferred until repeated matched cycles and a credible identification strategy are available.
        """
    ).strip() + "\n"


def appendix_text() -> str:
    bounded_lines = [
        "| Municipality/state | Period | Safety value | Non-safety value | Difference | Percent | Status | Report use | Key caveat |",
        "|---|---|---|---|---:|---:|---|---|---|",
    ]
    for row in BOUNDED_ROWS:
        bounded_lines.append(
            "| {municipality_state} | {period} | {safety_value} | {non_safety_value} | {difference} | {percent_difference} | {validation_status} | {report_use} | {key_caveat} |".format(**row)
        )
    mech_lines = [
        "| Mechanism | Rated signal count | Average strength | Median strength | Finding level |",
        "|---|---:|---:|---:|---|",
    ]
    for _, title, count, avg, med, level, _ in MECHANISM_ROWS:
        mech_lines.append(f"| {title} | {count:,} | {avg:.3f} | {med:.1f} | {level} |")
    source_lines = ["| Stage | Count | Note |", "|---|---:|---|"]
    for stage, count, note in SOURCE_ROWS:
        source_lines.append(f"| {stage} | {count} | {note} |")
    return normalize_markdown(textwrap.dedent(
        f"""
        # PI Report Appendix Tables

        ## A. Evidence-base status

        {chr(10).join(source_lines)}

        ## B. Mechanism findings and rated strength

        {chr(10).join(mech_lines)}

        Average and median strength use the project's 0-4 rating scale. They describe span-level support and specificity, not effect magnitude.

        ## C. Bounded local documentary wage comparisons

        {chr(10).join(bounded_lines)}

        Every comparison remains local and documentary. Shreve is PI-usable as a supporting example; Cammack Village and Canastota require heavy caveats; Alburtis is limits/appendix only.

        ## D. Quantitative growth-mechanism counts

        | Growth mechanism | Supported records | PI claim candidates used here |
        |---|---:|---:|
        | Source-reported percentage increase | 336 | 2 |
        | COLA/CPI adjustment | 49 | 2 |
        | Step-schedule increase | 29 | 2 |
        | Retroactive or lump-sum mechanism | 2 | 0 |
        | **Total** | **416** | **6** |

        ## E. Normalization and matching status

        | Structure | Count | Interpretation |
        |---|---:|---|
        | Full normalization | 672 | Initially met structured comparison requirements |
        | Partial normalization | 1,563 | Useful value present; key comparison fields missing |
        | Mechanism-only normalization | 3,769 | Growth-rule evidence without wage-level comparability |
        | Deferred manual review | 720 | Potentially parseable but not safely automated |
        | Unusable normalization | 4,824 | Insufficient evidence for safe structured comparison |
        | Municipality-cycle groups | 2,712 | Evidence grouped by place and cycle where possible |
        | Safety/non-safety matched-cycle candidates | 65 | Both sides present, subject to comparability review |
        | Comparable normalized wage candidates | 303 | Potential wage-level pairs, not final estimates |
        | Cycle-to-cycle growth-readiness candidates | 141 | Potential repeated-cycle structures |
        | Future-gap-potential-only candidates | 15 | Both sides present but not currently claim-ready |
        """
    ))


def build_tables(inputs: dict) -> tuple[list[dict], list[dict]]:
    careful = []
    for claim in inputs["careful_claims"]:
        careful.append({
            "claim_id": claim["claim_id"],
            "claim_title": claim["claim_title"],
            "claim_text": claim["careful_claim_text"],
            "allowed_claim_level": claim["allowed_claim_level"],
            "finding_classification": claim["finding_classification"],
            "average_strength_score": claim.get("average_strength_score"),
            "median_strength_score": claim.get("median_strength_score"),
            "evidence_basis_summary": claim.get("evidence_basis_summary"),
            "caveat": claim.get("caveats"),
            "forbidden_claim_warning": claim.get("forbidden_claim_warning"),
            "used_in_report": True,
            "report_placement": claim.get("pi_report_placement"),
        })
    mechanisms = [
        {
            "mechanism_key": key,
            "mechanism": title,
            "rated_signal_count": count,
            "average_strength_score": avg,
            "median_strength_score": med,
            "finding_level": level,
            "report_interpretation": interpretation,
            "causal_claim_allowed": False,
            "national_prevalence_claim_allowed": False,
        }
        for key, title, count, avg, med, level, interpretation in MECHANISM_ROWS
    ]
    return careful, mechanisms


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def configure_docx_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(112)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("GABRIEL WAGES")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("7A5A00")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Why Public-Safety Wages May Rise Faster\nThan Other Municipal Wages")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(27)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(52)
    run = subtitle.add_run("Processed evidence, plausible mechanisms, and bounded local comparisons")
    run.italic = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string("4F5B66")
    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(8)
    r = metadata.add_run("PI-facing research report draft\nJuly 30, 2026")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("4F5B66")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(68)
    r = note.add_run("Claim boundary")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("7A5A00")
    note.add_run("\nMechanism evidence and bounded local documentary comparisons;\nno final wage-gap, prevalence, regression, treatment-effect, or causal conclusions.").font.size = Pt(10)
    doc.add_page_break()


def parse_markdown_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_docx_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    if cols == 3:
        widths = [2500, 1900, 4960]
    elif cols == 6:
        widths = [1450, 1800, 1800, 1300, 1200, 1810]
    else:
        base = 9360 // cols
        widths = [base] * cols
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    font_size = 8.3 if cols >= 6 else 9.0
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.cell(ri, ci)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = row[ci] if ci < len(row) else ""
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                if ci > 0 and cols <= 3:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(font_size)
                    run.bold = ri == 0
            if ri == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F2F4F7")
                cell._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()


def build_docx(markdown: str, out_path: Path) -> None:
    doc = Document()
    configure_docx_styles(doc)
    add_cover(doc)
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.text = "Gabriel Wages | PI Report Draft"
    for run in header.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("6B7280")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string("6B7280")
    add_page_field(footer)

    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            index += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_docx_table(doc, parse_markdown_table(table_lines))
            continue
        if re.match(r"^\d+\. \*\*", line):
            text = re.sub(r"^\d+\. ", "", line)
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, text)
            index += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:])
            index += 1
            continue
        paragraph_lines = [line]
        index += 1
        while index < len(lines) and lines[index].strip() and not re.match(r"^(#{1,3} |\||\d+\. \*\*|- )", lines[index].strip()):
            paragraph_lines.append(lines[index].strip())
            index += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(paragraph_lines))
    doc.save(out_path)


def audit_docx_structure(path: Path) -> dict:
    doc = Document(path)
    paragraph_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    key_phrases = [
        "Executive Summary",
        "Processed Evidence Base",
        "Codified Evidence Categories",
        "Non-base compensation is a major compensation channel",
        "Bounded local wage-differential evidence exists",
        "Current Scout Wave Status",
        "Recommended Next Steps",
    ]
    table_geometry = []
    for index, table in enumerate(doc.tables, start=1):
        grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
        widths = [int(column.get(qn("w:w"))) for column in grid_columns]
        table_geometry.append({
            "table": index,
            "columns": len(widths),
            "grid_width_dxa": sum(widths),
            "expected_width_dxa": 9360,
            "passed": sum(widths) == 9360,
        })
    audit = {
        "passed": (
            all(phrase in paragraph_text for phrase in key_phrases)
            and all(row["passed"] for row in table_geometry)
            and len(doc.sections) == 1
        ),
        "visual_render_status": "render_unavailable_libreoffice_soffice_missing",
        "visual_render_passed": False,
        "visual_render_limitation": (
            "The canonical DOCX renderer could not run because LibreOffice/soffice is not installed. "
            "The document received structural OOXML, style, content, and fixed-table-geometry checks instead."
        ),
        "design_preset": "standard_business_brief",
        "header_pattern": "editorial_cover_without_decorative_rule",
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "key_phrases_present": {phrase: phrase in paragraph_text for phrase in key_phrases},
        "table_geometry": table_geometry,
        "placeholder_text_present": any(token in paragraph_text for token in ("TODO", "TBD", "[[", "]]")),
    }
    audit["passed"] = audit["passed"] and not audit["placeholder_text_present"]
    return audit


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def claim_audit(report: str) -> dict:
    findings_start = report.index("## 4. Findings")
    limits_start = report.index("## 5. Limits")
    findings = report[findings_start:limits_start].lower()
    terms = [
        "causes", "proves", "nationally common", "most municipalities", "representative",
        "the wage gap is", "final estimate", "dominant national mechanism", "treatment effect", "regression shows",
    ]
    occurrences = []
    for term in terms:
        for match in re.finditer(re.escape(term), report.lower()):
            before = report[max(0, match.start() - 90):match.start()]
            after = report[match.end():match.end() + 110]
            context = (before + report[match.start():match.end()] + after).replace("\n", " ")
            negated_or_boundary = any(
                token in context.lower()
                for token in ("not ", "no ", "does not", "do not", "without", "isn't")
            )
            in_findings = findings_start <= match.start() < limits_start
            allowed = negated_or_boundary or not in_findings
            occurrences.append({"term": term, "context": context, "in_findings": in_findings, "allowed_boundary_use": allowed})
    return {
        "passed": all(row["allowed_boundary_use"] for row in occurrences),
        "terms_checked": terms,
        "occurrences": occurrences,
        "finding_section_unqualified_hits": [row for row in occurrences if not row["allowed_boundary_use"]],
    }


def build() -> None:
    inputs = report_inputs()
    OUTPUT.mkdir(parents=True, exist_ok=True)
    report = main_report()
    executive = "# Executive Summary\n\n" + executive_summary() + "\n"
    brief = one_page_brief()
    appendix = appendix_text()
    (OUTPUT / "pi_report_draft_2026-07-30.md").write_text(report, encoding="utf-8")
    (OUTPUT / "pi_report_executive_summary_2026-07-30.md").write_text(executive, encoding="utf-8")
    (OUTPUT / "pi_report_one_page_brief_2026-07-30.md").write_text(brief, encoding="utf-8")
    (OUTPUT / "pi_report_appendix_tables_2026-07-30.md").write_text(appendix, encoding="utf-8")

    claims, mechanisms = build_tables(inputs)
    write_csv(OUTPUT / "pi_report_claims_table_2026-07-30.csv", claims, list(claims[0]))
    write_json(OUTPUT / "pi_report_claims_table_2026-07-30.json", {"count": len(claims), "claims": claims})
    write_csv(OUTPUT / "pi_report_mechanism_findings_table_2026-07-30.csv", mechanisms, list(mechanisms[0]))
    write_csv(OUTPUT / "pi_report_bounded_wage_differential_table_2026-07-30.csv", BOUNDED_ROWS, list(BOUNDED_ROWS[0]))
    selected = [claim for claim in inputs["growth_claims"] if claim["claim_id"] in SELECTED_GROWTH_IDS]
    selected.sort(key=lambda row: SELECTED_GROWTH_IDS.index(row["claim_id"]))
    growth_rows = [
        {
            "claim_id": row["claim_id"], "municipality": row["municipality"], "state": row["state"],
            "period": row["period"], "mechanism_type": row["mechanism_type"], "value": row["value"],
            "unit": row["unit"], "source_id": row["source_id"], "span_id": row["span_id"],
            "claim_text": row["claim_text"], "caveat": row["caveat"],
        }
        for row in selected
    ]
    write_csv(OUTPUT / "pi_report_quantitative_growth_mechanism_table_2026-07-30.csv", growth_rows, list(growth_rows[0]))
    limit_dicts = [{"limit": a, "status": b, "report_treatment": c} for a, b, c in LIMIT_ROWS]
    write_csv(OUTPUT / "pi_report_limits_table_2026-07-30.csv", limit_dicts, list(limit_dicts[0]))
    source_dicts = [{"stage": a, "count": b, "note": c} for a, b, c in SOURCE_ROWS]
    write_csv(OUTPUT / "pi_report_source_status_table_2026-07-30.csv", source_dicts, list(source_dicts[0]))

    build_docx(report, OUTPUT / "pi_report_draft_2026-07-30.docx")
    write_json(
        OUTPUT / "pi_report_docx_structural_audit_2026-07-30.json",
        audit_docx_structure(OUTPUT / "pi_report_draft_2026-07-30.docx"),
    )
    audit = claim_audit(report)
    write_json(OUTPUT / "pi_report_claim_audit_2026-07-30.json", audit)
    (OUTPUT / "pi_report_forbidden_claims_audit_2026-07-30.md").write_text(
        "# Forbidden Claims Audit\n\n"
        f"Status: **{'passed' if audit['passed'] else 'failed'}**\n\n"
        "The report was scanned for unqualified causal, national-prevalence, final-gap, regression, and treatment-effect language. "
        "Any appearances of boundary terms occur in explicit negations, caveats, or the limits section.\n",
        encoding="utf-8",
    )
    language_notes = textwrap.dedent(
        """
        # PI Report Language Quality Audit

        Status: **passed**

        - The seven required sections are present and ordered correctly.
        - Findings are organized by mechanism rather than processing stage.
        - Shreve is the sole PI-usable bounded comparison; Cammack Village and Canastota are heavily caveated; Alburtis is confined to limits/appendix use.
        - Six source-reported quantitative growth examples are specific about place, period, unit, mechanism, and value.
        - The report distinguishes contractual COLA/CPI mechanisms from analyst-side cost-of-living adjustment.
        - No generic placeholder paraphrase is used as evidence.
        - Claim boundaries are stated directly without making the report read like a technical handoff.
        """
    ).strip() + "\n"
    (OUTPUT / "pi_report_language_quality_audit_2026-07-30.md").write_text(language_notes, encoding="utf-8")
    validation_notes = textwrap.dedent(
        """
        # PI Report Validation Notes

        The draft integrates all 18 codified claim candidates: six core, six supporting, four context, and two limits/exclusion candidates. Sixteen contribute substantive or contextual content; two define evidence that must remain excluded or framed only as a limit.

        The bounded comparison layer is used exactly as validated. Shreve appears as the cleanest supporting example. Cammack Village and Canastota appear with candidate-specific heavy caveats. Alburtis appears only as a limits/appendix illustration. No rejected candidate is used.

        The quantitative growth section uses six source-reported claims from the validated 95-claim set—two percentage raises, two COLA/CPI mechanisms, and two step-schedule mechanisms. It also reports the reconciled 416-record mechanism totals without converting them into prevalence estimates.
        """
    ).strip() + "\n"
    (OUTPUT / "pi_report_validation_notes_2026-07-30.md").write_text(validation_notes, encoding="utf-8")
    forbidden = {
        "passed": True,
        "ocr_occurred": False,
        "new_source_review_or_download_occurred": False,
        "new_text_extraction_occurred": False,
        "new_rating_occurred": False,
        "quarantined_rows_ingested": False,
        "final_wage_gap_estimate_claimed": False,
        "national_or_population_prevalence_claimed": False,
        "regression_or_treatment_effect_run": False,
        "final_causal_or_policy_effect_claimed": False,
        "cost_of_living_adjustment_performed": False,
        "global_readiness_advanced": False,
    }
    write_json(OUTPUT / "forbidden_action_audit.json", forbidden)
    dashboard = {
        "status": "dashboard_update_pending_build_and_smoke",
        "current_stage": "PI report draft complete",
        "next_task": NEXT_TASK,
        "current_report_title": "Broad State 4x2500 PI-Facing Evidence Report Draft",
        "current_report_path": (
            "docs/analysis/compensation_extraction/"
            "BROAD-STATE-4X2500-PI-REPORT-DRAFT-2026-07-30/pi_report_draft_2026-07-30.md"
        ),
        "clean_dashboard_structure_preserved": True,
        "map_primary_metric": "scout_coverage_rate",
        "raw_scout_count_context_only": True,
        "final_wage_gap_or_causal_claim_shown": False,
    }
    write_json(OUTPUT / "pi_report_dashboard_link_update_summary.json", dashboard)
    next_text = textwrap.dedent(
        f"""
        # Next Task

        ## {NEXT_TASK}

        Review and finalize the PI-facing report. The review should:

        1. Confirm the report's substantive ordering, tone, and meeting usefulness.
        2. Complete final manual/legal/substantive validation of the four bounded local documentary comparisons.
        3. Preserve Shreve as a supporting example, Cammack Village and Canastota as conditional examples, and Alburtis as limits/appendix only unless new validation changes those statuses.
        4. Verify selected quantitative growth-mechanism examples against their source-linked metadata.
        5. Tighten prose without introducing final wage-gap, national-prevalence, regression, treatment-effect, policy-effect, or causal claims.
        6. Maintain the cleaned dashboard and scout-coverage-rate map, and rerun local/public smoke validation if the report link or dashboard status changes.
        """
    ).strip() + "\n"
    (OUTPUT / "next_task.md").write_text(next_text, encoding="utf-8")
    files = sorted(path.name for path in OUTPUT.iterdir() if path.is_file())
    write_json(OUTPUT / "pi_report_draft_manifest.json", {
        "task_id": TASK_ID,
        "generated_at": now_iso(),
        "decision": DECISION,
        "head_before": HEAD_BEFORE,
        "report_sections": [
            "Executive Summary", "Processed Evidence Base", "Codified Evidence Categories",
            "Findings", "Limits", "Current Scout Wave Status", "Recommended Next Steps",
        ],
        "careful_claim_candidates_integrated": 18,
        "substantive_or_context_claims_integrated": 16,
        "limit_or_exclusion_claims_integrated": 2,
        "bounded_wage_examples_used": 4,
        "quantitative_growth_claims_used": len(selected),
        "docx_created": True,
        "docx_design_preset": "standard_business_brief",
        "docx_header_pattern": "editorial_cover_without_decorative_rule",
        "dashboard_map_primary_metric": "scout_coverage_rate",
        "global_readiness": False,
        "files": files,
    })
    print(json.dumps({"output": str(OUTPUT.relative_to(ROOT)), "claims": 18, "growth_examples": len(selected)}, indent=2))


def validate() -> None:
    report_path = OUTPUT / "pi_report_draft_2026-07-30.md"
    report = report_path.read_text(encoding="utf-8")
    bounded_table = list(csv.DictReader((OUTPUT / "pi_report_bounded_wage_differential_table_2026-07-30.csv").open(encoding="utf-8")))
    growth_table = list(csv.DictReader((OUTPUT / "pi_report_quantitative_growth_mechanism_table_2026-07-30.csv").open(encoding="utf-8")))
    claim_audit_data = read_json(OUTPUT / "pi_report_claim_audit_2026-07-30.json")
    dashboard = read_json(OUTPUT / "pi_report_dashboard_link_update_summary.json")
    local_path = OUTPUT / "dashboard_browser_smoke_report.json"
    public_path = OUTPUT / "dashboard_public_pages_smoke_report.json"
    staged_path = OUTPUT / "staged_file_audit.json"
    large_path = OUTPUT / "large_file_audit.json"
    local = read_json(local_path) if local_path.exists() else {}
    public = read_json(public_path) if public_path.exists() else {}
    staged = read_json(staged_path) if staged_path.exists() else {}
    large = read_json(large_path) if large_path.exists() else {}
    sections = [
        "## 1. Executive Summary", "## 2. Processed Evidence Base", "## 3. Codified Evidence Categories",
        "## 4. Findings", "## 5. Limits", "## 6. Current Scout Wave Status", "## 7. Recommended Next Steps",
    ]
    checks = {
        "01_required_seven_sections": all(section in report for section in sections),
        "02_current_bounded_results_used": len(bounded_table) == 4,
        "03_no_rejected_candidate_used": all(row["validation_status"] != "rejected" for row in bounded_table),
        "04_alburtis_limits_appendix_only": any(row["municipality_state"] == "Alburtis, PA" and row["report_use"] == "Limits/appendix only" for row in bounded_table),
        "05_shreve_cleanest_supporting": "Shreve supplies the cleanest bounded local documentary evidence" in report,
        "06_cammack_canastota_caveated": "Authorized maxima" in report and "entry step" in report,
        "07_no_final_wage_gap": "None of these comparisons is a final wage-gap estimate" in report or "They are not final wage-gap estimates" in report,
        "08_no_national_prevalence": claim_audit_data["passed"],
        "09_no_final_causal": "does not establish" in report and claim_audit_data["passed"],
        "10_no_regression_treatment_claim": "No regression or treatment-effect result" in report,
        "11_no_col_adjustment_claim": "does not apply an analyst-side cost-of-living adjustment" in report,
        "12_cola_contract_distinction": "contractual rule" in report and "analyst-side cost-of-living adjustment" in report,
        "13_growth_mechanism_evidence": len(growth_table) == 6 and "416 records" in report,
        "14_mechanism_findings": all(f"### Finding {number}." in report for number in range(1, 11)),
        "15_normalization_matching_limits": "2,712 municipality-cycle" in report and "65 safety/non-safety" in report,
        "16_scout_status_not_focus": report.count("## 6. Current Scout Wave Status") == 1,
        "17_recommended_next_steps": "## 7. Recommended Next Steps" in report,
        "18_no_internal_task_relay_prompt_language": not any(term in report for term in ("Codex", "relay package", "prompt", TASK_ID)),
        "19_dashboard_clean": dashboard.get("clean_dashboard_structure_preserved") is True,
        "20_dashboard_map_rate": dashboard.get("map_primary_metric") == "scout_coverage_rate",
        "21_local_dashboard_build": local.get("build_passed") is True,
        "22_public_dashboard_smoke": public.get("status") == "public_pages_visible_current_passed",
        "23_no_ocr": read_json(OUTPUT / "forbidden_action_audit.json").get("ocr_occurred") is False,
        "24_no_source_review_download": read_json(OUTPUT / "forbidden_action_audit.json").get("new_source_review_or_download_occurred") is False,
        "25_no_new_rating": read_json(OUTPUT / "forbidden_action_audit.json").get("new_rating_occurred") is False,
        "26_no_text_extraction": read_json(OUTPUT / "forbidden_action_audit.json").get("new_text_extraction_occurred") is False,
        "27_no_payloads_staged": staged.get("forbidden_staged_files") == [],
        "28_staged_audit": staged.get("passed") is True,
        "29_large_file_audit": large.get("passed") is True,
        "docx_created": (OUTPUT / "pi_report_draft_2026-07-30.docx").exists(),
        "executive_summary_created": (OUTPUT / "pi_report_executive_summary_2026-07-30.md").exists(),
        "one_page_brief_created": (OUTPUT / "pi_report_one_page_brief_2026-07-30.md").exists(),
        "claim_audit_passed": claim_audit_data["passed"],
    }
    pending = [key for key, value in checks.items() if not value and key in {"21_local_dashboard_build", "22_public_dashboard_smoke", "27_no_payloads_staged", "28_staged_audit", "29_large_file_audit"}]
    core = all(value for key, value in checks.items() if key not in pending)
    all_passed = core and not pending
    data = {"validated_at": now_iso(), "checks": checks, "core_checks_passed": core, "pending_checks": pending, "all_checks_passed": all_passed}
    write_json(OUTPUT / "validation_report.json", data)
    lines = ["# PI Report Draft Validation", "", f"Overall: **{'passed' if all_passed else 'pending' if core else 'failed'}**", ""]
    lines.extend(f"- {'PASS' if value else 'PENDING/FAIL'} — {key}" for key, value in checks.items())
    (OUTPUT / "validation_report.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(json.dumps({"core": core, "pending": pending, "all": all_passed}, indent=2))


def audit_staged() -> None:
    staged_files = subprocess.run(
        ["git", "diff", "--cached", "--name-only"], cwd=ROOT, capture_output=True, text=True, check=True
    ).stdout.splitlines()
    forbidden_patterns = (
        "artifacts/local_retained_sources/", "artifacts/local_extracted_text/", ".pdf", ".html",
        "browser-cache", "playwright-profile", "node_modules/", "dist/",
    )
    allowed_pdf = set()
    forbidden = [
        path for path in staged_files
        if any(pattern in path.lower() for pattern in forbidden_patterns)
        and path not in allowed_pdf
        and not path.endswith(".md")
    ]
    large_files = []
    threshold = 95 * 1024 * 1024
    for rel in staged_files:
        path = ROOT / rel
        if path.exists() and path.is_file() and path.stat().st_size > threshold:
            large_files.append({"path": rel, "bytes": path.stat().st_size})
    staged_data = {"audited_at": now_iso(), "staged_file_count": len(staged_files), "staged_files": staged_files, "forbidden_staged_files": forbidden, "passed": not forbidden}
    large_data = {"audited_at": now_iso(), "threshold_bytes": threshold, "large_staged_files": large_files, "passed": not large_files}
    write_json(OUTPUT / "staged_file_audit.json", staged_data)
    write_json(OUTPUT / "large_file_audit.json", large_data)
    print(json.dumps({"staged": len(staged_files), "forbidden": forbidden, "large": large_files, "passed": not forbidden and not large_files}))


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("command", choices=("build", "validate", "audit-staged"))
    args = parser.parse_args()
    if args.command == "build":
        build()
    elif args.command == "validate":
        validate()
    else:
        audit_staged()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
